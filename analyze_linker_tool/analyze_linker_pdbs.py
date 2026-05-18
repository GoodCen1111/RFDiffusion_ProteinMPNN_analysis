import csv
import glob
import math
import os
from collections import defaultdict


ROOT_DIR = os.path.abspath(os.path.dirname(__file__))
OUTPUT_CSV = os.path.join(ROOT_DIR, "linker_analysis_ranking_strict_with_B.csv")
OUTPUT_DOCX = os.path.join(ROOT_DIR, "linker_analysis_report_strict_zh_with_B.docx")


def parse_pdb_atoms(pdb_path):
    atoms = []
    with open(pdb_path, "r", encoding="utf-8", errors="ignore") as handle:
        for line in handle:
            if not line.startswith("ATOM"):
                continue
            atom_name = line[12:16].strip()
            res_name = line[17:20].strip()
            chain_id = line[21].strip()
            try:
                res_seq = int(line[22:26])
                x = float(line[30:38])
                y = float(line[38:46])
                z = float(line[46:54])
            except ValueError:
                continue
            atoms.append(
                {
                    "atom_name": atom_name,
                    "res_name": res_name,
                    "chain_id": chain_id,
                    "res_seq": res_seq,
                    "xyz": (x, y, z),
                }
            )
    return atoms


def distance(a, b):
    dx = a[0] - b[0]
    dy = a[1] - b[1]
    dz = a[2] - b[2]
    return math.sqrt(dx * dx + dy * dy + dz * dz)


def dihedral(p1, p2, p3, p4):
    def vsub(a, b):
        return (a[0] - b[0], a[1] - b[1], a[2] - b[2])

    def dot(a, b):
        return a[0] * b[0] + a[1] * b[1] + a[2] * b[2]

    def cross(a, b):
        return (
            a[1] * b[2] - a[2] * b[1],
            a[2] * b[0] - a[0] * b[2],
            a[0] * b[1] - a[1] * b[0],
        )

    def norm(a):
        return math.sqrt(dot(a, a))

    b1 = vsub(p2, p1)
    b2 = vsub(p3, p2)
    b3 = vsub(p4, p3)

    n1 = cross(b1, b2)
    n2 = cross(b2, b3)
    b2n = norm(b2)
    if b2n == 0:
        return None
    n1n = norm(n1)
    n2n = norm(n2)
    if n1n == 0 or n2n == 0:
        return None

    n1u = (n1[0] / n1n, n1[1] / n1n, n1[2] / n1n)
    n2u = (n2[0] / n2n, n2[1] / n2n, n2[2] / n2n)
    b2u = (b2[0] / b2n, b2[1] / b2n, b2[2] / b2n)
    m1 = cross(n1u, b2u)
    x = dot(n1u, n2u)
    y = dot(m1, n2u)
    angle = math.degrees(math.atan2(y, x))
    return angle


def radius_of_gyration(coords):
    if not coords:
        return None
    cx = sum(c[0] for c in coords) / len(coords)
    cy = sum(c[1] for c in coords) / len(coords)
    cz = sum(c[2] for c in coords) / len(coords)
    rg2 = sum((c[0] - cx) ** 2 + (c[1] - cy) ** 2 + (c[2] - cz) ** 2 for c in coords) / len(coords)
    return math.sqrt(rg2)


def detect_linker_range(chain_res):
    # Longest continuous GLY run on chain A is treated as linker.
    best = None
    i = 0
    while i < len(chain_res):
        j = i
        while j < len(chain_res) and chain_res[j][1] == "GLY":
            j += 1
        run_len = j - i
        if run_len >= 10:
            start = chain_res[i][0]
            end = chain_res[j - 1][0]
            if best is None or run_len > best[2]:
                best = (start, end, run_len)
        i = j if j > i else i + 1
    return best


def get_chain_residues(atoms, chain_id="A"):
    seen = set()
    residues = []
    for a in atoms:
        if a["chain_id"] != chain_id:
            continue
        key = a["res_seq"]
        if key in seen:
            continue
        seen.add(key)
        residues.append((a["res_seq"], a["res_name"]))
    residues.sort(key=lambda x: x[0])
    return residues


def compute_metrics(pdb_path):
    atoms = parse_pdb_atoms(pdb_path)
    if not atoms:
        return None

    chain_a_res = get_chain_residues(atoms, "A")
    linker = detect_linker_range(chain_a_res)
    if linker is None:
        return None
    link_start, link_end, link_len = linker

    atom_by_chain_res = defaultdict(dict)
    for a in atoms:
        atom_by_chain_res[(a["chain_id"], a["res_seq"])][a["atom_name"]] = a["xyz"]

    link_ca = []
    for r in range(link_start, link_end + 1):
        ca = atom_by_chain_res.get(("A", r), {}).get("CA")
        if ca is not None:
            link_ca.append((r, ca))

    ca_step_deviation = None
    ca_step_bad = 0
    ca_step_count = 0
    if len(link_ca) > 1:
        steps = []
        for i in range(len(link_ca) - 1):
            d = distance(link_ca[i][1], link_ca[i + 1][1])
            steps.append(abs(d - 3.8))
            ca_step_count += 1
            if d < 2.8 or d > 4.5:
                ca_step_bad += 1
        ca_step_deviation = sum(steps) / len(steps)

    rg = radius_of_gyration([p[1] for p in link_ca]) if link_ca else None

    n_xyz = atom_by_chain_res.get(("A", link_start), {}).get("N")
    c_before = atom_by_chain_res.get(("A", link_start - 1), {}).get("C")
    c_xyz = atom_by_chain_res.get(("A", link_end), {}).get("C")
    n_after = atom_by_chain_res.get(("A", link_end + 1), {}).get("N")

    left_peptide = distance(c_before, n_xyz) if (c_before is not None and n_xyz is not None) else None
    right_peptide = distance(c_xyz, n_after) if (c_xyz is not None and n_after is not None) else None
    cn_pairs = []
    for r in range(link_start - 1, link_end + 1):
        c_i = atom_by_chain_res.get(("A", r), {}).get("C")
        n_i1 = atom_by_chain_res.get(("A", r + 1), {}).get("N")
        if c_i is not None and n_i1 is not None:
            cn_pairs.append((r, r + 1, distance(c_i, n_i1)))

    cn_compressed = 0
    cn_stretched = 0
    for _, _, d in cn_pairs:
        if d < 1.15:
            cn_compressed += 1
        if d > 2.2:
            cn_stretched += 1

    omega_total = 0
    omega_cis_like = 0
    omega_bad = 0
    omega_values = []
    for r in range(link_start - 1, link_end + 1):
        ca_i = atom_by_chain_res.get(("A", r), {}).get("CA")
        c_i = atom_by_chain_res.get(("A", r), {}).get("C")
        n_i1 = atom_by_chain_res.get(("A", r + 1), {}).get("N")
        ca_i1 = atom_by_chain_res.get(("A", r + 1), {}).get("CA")
        if None in (ca_i, c_i, n_i1, ca_i1):
            continue
        w = dihedral(ca_i, c_i, n_i1, ca_i1)
        if w is None:
            continue
        omega_total += 1
        omega_values.append(w)
        # trans ~180, cis ~0
        if abs(w) < 30.0:
            omega_cis_like += 1
        if abs(abs(w) - 180.0) > 60.0:
            omega_bad += 1

    # Clash counting between linker atoms and non-linker atoms on chain A.
    linker_atoms = []
    non_linker_atoms = []
    b_chain_atoms = []
    for a in atoms:
        if a["chain_id"] == "A":
            if link_start <= a["res_seq"] <= link_end:
                linker_atoms.append(a)
            else:
                non_linker_atoms.append(a)
        elif a["chain_id"] == "B":
            b_chain_atoms.append(a)

    severe_clash = 0
    mild_clash = 0
    try:
        from scipy.spatial import cKDTree  # type: ignore

        points = [a["xyz"] for a in non_linker_atoms]
        tree = cKDTree(points) if points else None
        for la in linker_atoms:
            if tree is None:
                break
            idxs = tree.query_ball_point(la["xyz"], 2.4)
            for idx in idxs:
                oa = non_linker_atoms[idx]
                # Exclude near-sequence bonded neighborhood (1-3 exclusion)
                if abs(la["res_seq"] - oa["res_seq"]) <= 2:
                    continue
                d = distance(la["xyz"], oa["xyz"])
                if d < 2.0:
                    severe_clash += 1
                elif d < 2.4:
                    mild_clash += 1
    except Exception:
        # Fallback without scipy
        for la in linker_atoms:
            for oa in non_linker_atoms:
                if abs(la["res_seq"] - oa["res_seq"]) <= 2:
                    continue
                d = distance(la["xyz"], oa["xyz"])
                if d < 2.0:
                    severe_clash += 1
                elif d < 2.4:
                    mild_clash += 1

    # Clash between linker (chain A) and chain B atoms.
    severe_clash_b = 0
    mild_clash_b = 0
    try:
        from scipy.spatial import cKDTree  # type: ignore

        b_points = [a["xyz"] for a in b_chain_atoms]
        tree_b = cKDTree(b_points) if b_points else None
        for la in linker_atoms:
            if tree_b is None:
                break
            idxs = tree_b.query_ball_point(la["xyz"], 2.4)
            for idx in idxs:
                ba = b_chain_atoms[idx]
                d = distance(la["xyz"], ba["xyz"])
                if d < 2.0:
                    severe_clash_b += 1
                elif d < 2.4:
                    mild_clash_b += 1
    except Exception:
        for la in linker_atoms:
            for ba in b_chain_atoms:
                d = distance(la["xyz"], ba["xyz"])
                if d < 2.0:
                    severe_clash_b += 1
                elif d < 2.4:
                    mild_clash_b += 1

    # End-to-end consistency for polymer-like linker.
    end_to_end = None
    contour = None
    ext_ratio = None
    if link_ca:
        end_to_end = distance(link_ca[0][1], link_ca[-1][1])
        contour = 3.8 * (len(link_ca) - 1)
        if contour > 0:
            ext_ratio = end_to_end / contour

    # Score components (higher is better)
    score = 100.0
    reasons = []

    if ca_step_deviation is not None:
        penalty = min(18.0, ca_step_deviation * 20.0)
        score -= penalty
        reasons.append(f"CA-CA 平均偏差(相对3.8A)：{ca_step_deviation:.3f}A")
    else:
        score -= 20.0
        reasons.append("Linker 缺少 CA 数据")

    if ca_step_count > 0:
        bad_ratio = ca_step_bad / ca_step_count
        score -= min(20.0, bad_ratio * 40.0)
        reasons.append(f"CA 步长异常(<2.8 或 >4.5A)：{ca_step_bad}/{ca_step_count}")

    for side_name, val in [("左端", left_peptide), ("右端", right_peptide)]:
        if val is None:
            score -= 10.0
            reasons.append(f"{side_name}边界肽键缺失")
            continue
        if val < 1.15:
            score -= 12.0
            reasons.append(f"{side_name} C-N={val:.3f}A（异常压缩）")
        elif val > 2.2:
            score -= 25.0
            reasons.append(f"{side_name} C-N={val:.3f}A（断裂/严重拉伸）")
        else:
            dev = abs(val - 1.33)
            score -= min(6.0, dev * 10.0)
            reasons.append(f"{side_name} C-N={val:.3f}A（正常范围）")

    # Penalize all C(i)-N(i+1) around linker
    score -= min(18.0, cn_compressed * 4.0 + cn_stretched * 6.0)
    reasons.append(
        f"全段 C(i)-N(i+1) 异常：压缩(<1.15A)={cn_compressed}，拉伸(>2.2A)={cn_stretched}，总对数={len(cn_pairs)}"
    )

    # Omega penalties
    score -= min(20.0, omega_cis_like * 4.0 + omega_bad * 2.0)
    reasons.append(
        f"ω 二面角：cis-like(|ω|<30°)={omega_cis_like}，明显偏离trans={omega_bad}，总计={omega_total}"
    )

    score -= min(20.0, severe_clash * 0.9 + mild_clash * 0.2)
    reasons.append(f"空间冲突（排除1-3邻近）：严重<2.0A={severe_clash}，轻度<2.4A={mild_clash}")

    # B-chain interface clashes: strong penalty because this is user-requested key criterion.
    score -= min(30.0, severe_clash_b * 1.0 + mild_clash_b * 0.25)
    reasons.append(f"与B链冲突：严重<2.0A={severe_clash_b}，轻度<2.4A={mild_clash_b}")

    if ext_ratio is not None:
        # Flexible linker usually around 0.4-0.6
        if ext_ratio < 0.4:
            score -= min(12.0, (0.4 - ext_ratio) * 40.0)
            reasons.append(f"L/L0={ext_ratio:.3f}（偏塌缩，低于0.4）")
        elif ext_ratio > 0.6:
            score -= min(12.0, (ext_ratio - 0.6) * 40.0)
            reasons.append(f"L/L0={ext_ratio:.3f}（偏拉伸，高于0.6）")
        else:
            reasons.append(f"L/L0={ext_ratio:.3f}（柔性区间内）")

    if rg is not None:
        reasons.append(f"Linker CA 回转半径 Rg={rg:.3f}A")

    score = max(0.0, min(100.0, score))

    return {
        "file": os.path.relpath(pdb_path, ROOT_DIR),
        "basename": os.path.basename(pdb_path),
        "linker_start": link_start,
        "linker_end": link_end,
        "linker_len": link_len,
        "ca_step_dev": ca_step_deviation,
        "ca_step_bad": ca_step_bad,
        "ca_step_total": ca_step_count,
        "left_cn": left_peptide,
        "right_cn": right_peptide,
        "cn_compressed": cn_compressed,
        "cn_stretched": cn_stretched,
        "omega_total": omega_total,
        "omega_cis_like": omega_cis_like,
        "omega_bad": omega_bad,
        "severe_clash": severe_clash,
        "mild_clash": mild_clash,
        "severe_clash_b": severe_clash_b,
        "mild_clash_b": mild_clash_b,
        "rg": rg,
        "ext_ratio": ext_ratio,
        "score": score,
        "reasons": " | ".join(reasons),
    }


def write_csv(results):
    fields = [
        "rank",
        "score",
        "file",
        "linker_start",
        "linker_end",
        "linker_len",
        "ca_step_dev",
        "ca_step_bad",
        "ca_step_total",
        "left_cn",
        "right_cn",
        "cn_compressed",
        "cn_stretched",
        "omega_total",
        "omega_cis_like",
        "omega_bad",
        "severe_clash",
        "mild_clash",
        "severe_clash_b",
        "mild_clash_b",
        "rg",
        "ext_ratio",
        "reasons",
    ]
    with open(OUTPUT_CSV, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for i, r in enumerate(results, start=1):
            row = {"rank": i}
            for k in fields:
                if k == "rank":
                    continue
                row[k] = r.get(k)
            writer.writerow(row)


def write_docx(results):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed. Please install with: pip install python-docx") from exc

    doc = Document()
    doc.add_heading("RFdiffusion Linker PDB 严格标准分析报告（中文）", level=1)
    doc.add_paragraph(
        "数据集：0-99 共100个PDB。目标：A链中连续GLY linker（作为未知氨基酸占位，保留不修改）。"
        "本报告采用更严格的几何与冲突标准评分，并按分数从高到低排序。"
    )

    doc.add_heading("评分标准（严格版）", level=2)
    doc.add_paragraph(
        "1) 肽键 C(i)-N(i+1)：理想约1.33A，<1.15A记为异常压缩，>2.2A记为断裂/严重拉伸；\n"
        "2) CA(i)-CA(i+1)：步长合理区间2.8-4.5A，统计异常比例；\n"
        "3) 肽平面ω（CA-C-N-CA）：统计cis-like与偏离trans异常；\n"
        "4) 空间冲突：A链内部冲突（排除1-3邻近）与A-linker对B链冲突，分别统计<2.0A与<2.4A；\n"
        "5) Linker形状：Rg与L/L0（柔性linker常见0.4-0.6）。"
    )

    doc.add_heading("排名结果（高分到低分）", level=2)
    for i, r in enumerate(results, start=1):
        p = doc.add_paragraph()
        p.add_run(f"#{i} | 分数 {r['score']:.2f} | {r['basename']}").bold = True
        doc.add_paragraph(
            f"Linker区间：A:{r['linker_start']}-{r['linker_end']}（长度={r['linker_len']}）\n"
            f"CA偏差={r['ca_step_dev']:.3f}A；CA异常步长={r['ca_step_bad']}/{r['ca_step_total']}；"
            f"C-N(左/右)={r['left_cn']:.3f}/{r['right_cn']:.3f}A；"
            f"C-N异常(压缩/拉伸)={r['cn_compressed']}/{r['cn_stretched']}；"
            f"ω(cis-like/异常/总数)={r['omega_cis_like']}/{r['omega_bad']}/{r['omega_total']}；"
            f"A链内冲突(严重/轻度)={r['severe_clash']}/{r['mild_clash']}；"
            f"与B链冲突(严重/轻度)={r['severe_clash_b']}/{r['mild_clash_b']}；"
            f"Rg={r['rg']:.3f}A；L/L0={r['ext_ratio']:.3f}\n"
            f"原因：{r['reasons']}"
        )

    doc.save(OUTPUT_DOCX)


def main():
    pdb_files = sorted(glob.glob(os.path.join(ROOT_DIR, "**", "*.pdb"), recursive=True))
    # Keep canonical 0-99 set, ignore accidental duplicate names like 90[1].
    canonical = []
    seen_idx = set()
    for p in pdb_files:
        base = os.path.basename(p)
        # Expected name: BJ_linker_IT_design_<idx>.pdb
        parts = base.replace(".pdb", "").split("_")
        try:
            idx = int(parts[-1])
        except ValueError:
            continue
        if 0 <= idx <= 99 and idx not in seen_idx:
            seen_idx.add(idx)
            canonical.append((idx, p))

    canonical.sort(key=lambda x: x[0])

    results = []
    for idx, pdb in canonical:
        m = compute_metrics(pdb)
        if m is None:
            continue
        m["design_index"] = idx
        results.append(m)

    results.sort(key=lambda x: x["score"], reverse=True)
    write_csv(results)
    write_docx(results)
    print(f"Analyzed {len(results)} models.")
    print(f"CSV:  {OUTPUT_CSV}")
    print(f"DOCX: {OUTPUT_DOCX}")
    if results:
        print("Top 10:")
        for i, r in enumerate(results[:10], start=1):
            print(f"{i:>2}. {r['basename']}  score={r['score']:.2f}  linker=A:{r['linker_start']}-{r['linker_end']}")


if __name__ == "__main__":
    main()
